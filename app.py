from flask import Flask, render_template, request, send_file
from docx import Document
import PyPDF2
from google import genai
import os

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# =========================
# Gemini Client
# =========================
# =========================
# Gemini Client (Environment Variable)
# =========================
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise ValueError("❌ GEMINI_API_KEY not found in environment variables!")

client = genai.Client(api_key=GEMINI_API_KEY)

# =========================
# Functions
# =========================
def summarize_text(text):
    prompt = f"Please summarize the following text clearly:\n\n{text}"
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt
    )
    return response.text if response.text else "No summary generated."

def read_word(path):
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs])

def read_pdf(path):
    text = ""
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def save_to_word(summary):
    file_path = "summary.docx"
    doc = Document()
    doc.add_heading("Automatic Summary", level=1)
    doc.add_paragraph(summary)
    doc.save(file_path)
    return file_path

# =========================
# Routes
# =========================
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        file = request.files.get("file")
        if not file:
            return "❌ No file uploaded!", 400

        path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(path)

        # اقرأ النص حسب نوع الملف
        if file.filename.lower().endswith(".pdf"):
            text = read_pdf(path)
        elif file.filename.lower().endswith(".docx") or file.filename.lower().endswith(".doc"):
            text = read_word(path)
        else:
            return "❌ Unsupported file type! Use PDF or Word.", 400

        # تلخيص النص
        summary = summarize_text(text)

        # حفظ الملخص وارساله كملف Word
        word_file = save_to_word(summary)
        return send_file(word_file, as_attachment=True)

    return render_template("index.html")

# =========================
# Run App
# =========================
if __name__ == "__main__":
    # Local dev mode
    app.run(debug=True)
